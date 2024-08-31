# Date: 20.08.2024.
# Name: Ruslan Yarmak.
# Contacts: (Phone number: +375297242242, Telegram: @ruslanyarmak).

from PyQt5 import QtCore, QtWidgets
from PyQt5.QtWidgets import QTableWidget, QMessageBox
from functools import partial
import sqlite3
import sys
import subprocess
import os
import platform
from Sergey.client_card import Ui_Client_Add
from Sergey.authorization import Ui_Authorization
from Sergey.list_of_manager import Ui_listManager
from Maksim.warehouse_dialog import AddWarehouseDialog
from Maksim.product_dialog import AddProductDialog
from Alex.test import DialogWindow
from Ruslan.data_processing_class import DataProcessing


class MainDialog(QtWidgets.QDialog):
    def __init__(self, user=None, manager_id=None, pattern_doc=None, file_db='My/store_database.db'):
        super().__init__()
        self.manager_id = manager_id
        self.user = user
        self.pattern_doc = pattern_doc
        self.con = sqlite3.connect(file_db)
        self.table_widgets = []  # Ссылки на виджеты таблиц.
        self.lineedit_list = []  # Ссылки на лайн эдиты.
        self.table_names = ['Transactions_history', 'Customers', 'Stock', 'Products']

        self.setFixedSize(1200, 940)

        self.tabWidget = QtWidgets.QTabWidget(self)
        self.tabWidget.setGeometry(QtCore.QRect(0, 0, 1200, 940))

# -------------Вкладка-Операций----------------------------------------------------------------------------------------
        self.tab = QtWidgets.QWidget()

        self.tableWidget = QtWidgets.QTableWidget(self.tab)
        self.tableWidget.setGeometry(QtCore.QRect(10, 50, 1175, 850))
        self.tableWidget.setEditTriggers(QTableWidget.NoEditTriggers)
        self.table_widgets.append(self.tableWidget)

        self.new_transaction = QtWidgets.QPushButton(self.tab)
        self.new_transaction.setGeometry(QtCore.QRect(1040, 10, 131, 31))

        self.upload_document = QtWidgets.QPushButton(self.tab)
        self.upload_document.setGeometry(QtCore.QRect(460, 10, 131, 31))

        if user == 'super':
            self.add_word = QtWidgets.QPushButton(self.tab)
            self.add_word.setGeometry(QtCore.QRect(600, 10, 70, 31))
            self.add_word.setText('Add Word')

            self.add_excel = QtWidgets.QPushButton(self.tab)
            self.add_excel.setGeometry(QtCore.QRect(680, 10, 70, 31))
            self.add_excel.setText('Add Excel')

            self.add_manager = QtWidgets.QPushButton(self.tab)
            self.add_manager.setGeometry(QtCore.QRect(760, 10, 131, 31))

            self.managers = QtWidgets.QPushButton(self.tab)
            self.managers.setGeometry(QtCore.QRect(900, 10, 131, 31))
            self.managers.setText('Managers')

        self.comboBox_2 = QtWidgets.QComboBox(self.tab)
        self.comboBox_2.setGeometry(QtCore.QRect(345, 16, 105, 21))
        #self.add_warehouses_to_combobox()

        self.lineEdit = QtWidgets.QLineEdit(self.tab)
        self.lineEdit.setGeometry(QtCore.QRect(180, 16, 153, 21))
        self.lineEdit.setPlaceholderText('Поиск...')
        self.lineedit_list.append(self.lineEdit)

        self.tabWidget.addTab(self.tab, "")

# -------------Вкладка-Складов----------------------------------------------------------------------------------------
        self.tab_2 = QtWidgets.QWidget()

        self.tableWidget_2 = QtWidgets.QTableWidget(self.tab_2)
        self.tableWidget_2.setGeometry(QtCore.QRect(10, 50, 1175, 850))
        self.tableWidget_2.setEditTriggers(QTableWidget.NoEditTriggers)
        self.table_widgets.append(self.tableWidget_2)

        self.lineEdit_2 = QtWidgets.QLineEdit(self.tab_2)
        self.lineEdit_2.setGeometry(QtCore.QRect(180, 16, 153, 21))
        self.lineEdit_2.setPlaceholderText('Поиск...')
        self.lineedit_list.append(self.lineEdit_2)

        self.add_client = QtWidgets.QPushButton(self.tab_2)
        self.add_client.setGeometry(QtCore.QRect(1040, 10, 131, 31))

        self.tabWidget.addTab(self.tab_2, "")

# -------------Вкладка-Клиентов----------------------------------------------------------------------------------------
        self.tab_3 = QtWidgets.QWidget()

        self.tableWidget_3 = QtWidgets.QTableWidget(self.tab_3)
        self.tableWidget_3.setGeometry(QtCore.QRect(10, 50, 1175, 850))
        self.tableWidget_3.setEditTriggers(QTableWidget.NoEditTriggers)
        self.table_widgets.append(self.tableWidget_3)

        self.lineEdit_3 = QtWidgets.QLineEdit(self.tab_3)
        self.lineEdit_3.setGeometry(QtCore.QRect(180, 16, 153, 21))
        self.lineEdit_3.setPlaceholderText('Поиск...')
        self.lineedit_list.append(self.lineEdit_3)

        self.add_warehouse = QtWidgets.QPushButton(self.tab_3)
        self.add_warehouse.setGeometry(QtCore.QRect(1040, 10, 131, 31))

        self.edit_warehouse = QtWidgets.QPushButton(self.tab_3)
        self.edit_warehouse.setGeometry(QtCore.QRect(900, 10, 131, 31))

        self.comboBox = QtWidgets.QComboBox(self.tab_3)
        self.comboBox.setGeometry(QtCore.QRect(10, 16, 161, 21))
        self.add_warehouses_to_combobox()

        self.tabWidget.addTab(self.tab_3, "")

# -------------Вкладка-Продуктов----------------------------------------------------------------------------------------
        self.tab_4 = QtWidgets.QWidget()

        self.tableWidget_4 = QtWidgets.QTableWidget(self.tab_4)
        self.tableWidget_4.setGeometry(QtCore.QRect(10, 50, 1175, 850))
        self.tableWidget_4.setEditTriggers(QTableWidget.NoEditTriggers)
        self.table_widgets.append(self.tableWidget_4)

        self.lineEdit_4 = QtWidgets.QLineEdit(self.tab_4)
        self.lineEdit_4.setGeometry(QtCore.QRect(180, 16, 153, 21))
        self.lineEdit_4.setPlaceholderText('Поиск...')
        self.lineedit_list.append(self.lineEdit_4)

        self.add_product = QtWidgets.QPushButton(self.tab_4)
        self.add_product.setGeometry(QtCore.QRect(1040, 10, 131, 31))

        self.tabWidget.addTab(self.tab_4, "")
# --------------Сигналы-------------------------------------------------------------------------------------------------
        self.tabWidget.currentChanged.connect(self.insert_data_into_table)
        self.comboBox.activated.connect(self.insert_data_into_table)
        self.lineEdit.textChanged.connect(self.insert_data_into_table)
        self.lineEdit_2.textChanged.connect(self.insert_data_into_table)
        self.lineEdit_3.textChanged.connect(self.insert_data_into_table)
        self.lineEdit_4.textChanged.connect(self.insert_data_into_table)
        self.add_client.clicked.connect(self.open_client_card)
        self.tableWidget_2.cellDoubleClicked.connect(self.open_client_card)
        self.add_product.clicked.connect(self.open_product_card)
        self.tableWidget_4.cellDoubleClicked.connect(self.open_product_card)
        self.add_warehouse.clicked.connect(self.open_warehouse_card)
        self.edit_warehouse.clicked.connect(partial(self.open_warehouse_card, True))
        self.new_transaction.clicked.connect(self.open_transaction_card)
        self.tableWidget.cellDoubleClicked.connect(self.open_transaction_card)
        if user == 'super':
            self.add_manager.clicked.connect(self.open_manager_card)
            self.add_word.clicked.connect(self.add_word_doc)
            self.managers.clicked.connect(self.open_managers_list)

        self.announcement = QMessageBox()
        self.announcement.setWindowTitle('Info')

        self.retranslateUi()

        self.insert_data_into_table()

    def retranslateUi(self):
        _translate = QtCore.QCoreApplication.translate
        self.setWindowTitle(_translate("Dialog", "CRM Application"))
        self.new_transaction.setText(_translate("Dialog", "New transaction"))
        self.upload_document.setText(_translate("Dialog", "Upload document"))
        if self.user == 'super':
            self.add_word.setText(_translate("Dialog", "Add Word"))
            self.add_manager.setText(_translate("Dialog", "Add manager"))
            self.managers.setText(_translate("Dialog", "Managers"))
        self.tabWidget.setTabText(self.tabWidget.indexOf(self.tab), _translate("Dialog", "Operations"))
        self.add_client.setText(_translate("Dialog", "Add client"))
        self.tabWidget.setTabText(self.tabWidget.indexOf(self.tab_2), _translate("Dialog", "Сlients"))
        self.add_warehouse.setText(_translate("Dialog", "Add warehouse"))
        self.edit_warehouse.setText(_translate("Dialog", "Edit warehouse"))
        self.tabWidget.setTabText(self.tabWidget.indexOf(self.tab_3), _translate("Dialog", "Warehouses"))
        self.add_product.setText(_translate("Dialog", "Add product"))
        self.tabWidget.setTabText(self.tabWidget.indexOf(self.tab_4), _translate("Dialog", "Products"))

    def insert_data_into_table(self) -> None:
        """Функция заполняет таблицу данными"""
        tab_index = self.tabWidget.currentIndex()  # Индекс вкладки.
        search_txt = self.lineedit_list[tab_index].text().lower()  # Текс в поисковой строке.
        table = self.table_widgets[tab_index]  # Выбор tableWidget.
        table.setRowCount(0)  # Обнуление таблицы.
        index_warehouse = self.comboBox.currentIndex()  # Индекс склада в combobox.
        self.edit_warehouse.setEnabled(True)
        if index_warehouse == 0:
            self.edit_warehouse.setEnabled(False)
        last_index = 0
        with self.con:
            resp = self.con.execute(f"Pragma table_info ({self.table_names[tab_index]})").fetchall()
            columns_list = [i[1] for i in resp]
            table_data = self.con.execute(f"SELECT * FROM {self.table_names[tab_index]}").fetchall()
            if tab_index == 2 and index_warehouse == 0:
                columns_list, table_data = self.count_quantity_products()
            elif tab_index == 2 and index_warehouse > 0:
                resp = self.con.execute('''SELECT product_id,
                                                  product_name,
                                                  warehouse_id,
                                                  sku,
                                                  unit_of_measurement,
                                                  quantity,  
                                                  expiration_date
                                             FROM Products
                                             JOIN Stock
                                               ON Products.id = Stock.product_id
                                            ORDER BY expiration_date
                ''')
                table_data = resp.fetchall()
                columns_list = [el[0] for el in resp.description]
            rows_list = [str(i[0]) for i in table_data]
            table.setColumnCount(len(columns_list))
            table.setHorizontalHeaderLabels(columns_list)
            table.setVerticalHeaderLabels(rows_list)
            for i in range(len(table_data)):
                if index_warehouse != table_data[i][2] and tab_index == 2 and index_warehouse > 0:
                    continue
                if search_txt not in str(table_data[i][1]).lower():
                    continue
                table.insertRow(last_index)
                for j in range(len(table_data[i])):
                    item = QtWidgets.QTableWidgetItem(str(table_data[i][j]))
                    table.setItem(last_index, j, item)
                last_index += 1

    def add_warehouses_to_combobox(self):
        """Добавляет имена складов в ComboBox"""
        self.comboBox.clear()
        self.comboBox.addItems(['Stock'])
        warehouse_names = [el[0] for el in self.con.execute("SELECT warehouse_name FROM Warehouses").fetchall()]
        self.comboBox.addItems(warehouse_names)

    def get_string_values(self, row: int) -> list:
        """Функция принимает индекс выбранной ячейки в TableWidget и возвращает все значения строки"""
        column_count = self.table_widgets[self.tabWidget.currentIndex()].columnCount()
        values = []
        for column in range(column_count):
            item = self.table_widgets[self.tabWidget.currentIndex()].item(row, column)
            if item is not None:
                values.append(item.text())
            else:
                values.append('')
        return values

    def open_client_card(self, arg):
        """Открывает карточку клиента и передает первым аргументом False (если функция вызвана по нажатию PushButton),
        либо список всех значений выбранной строки (если функция вызвана двойным нажатием по TableWidget)."""
        if arg is not False:
            arg = self.get_string_values(arg)
        client_card_window = Ui_Client_Add(arg)
        client_card_window.exec_()
        self.insert_data_into_table()

    def open_manager_card(self):
        """Открывает карточку менеджера"""
        manager_card_window = Ui_Authorization()
        manager_card_window.exec_()

    def open_warehouse_card(self, arg: bool):
        """Открывает карточку склада и передает False (если функция вызвана по нажатию PushButton),
        либо кортеж всех значений из текущего склада выбранного в ComboBox."""
        warehouse = self.comboBox.currentText()
        if arg:
            arg = self.con.execute(f"SELECT * FROM Warehouses WHERE warehouse_name = '{warehouse}'").fetchall()[0]
        warehouse_card_window = AddWarehouseDialog(arg)
        warehouse_card_window.exec_()
        self.add_warehouses_to_combobox()
        self.insert_data_into_table()

    def open_product_card(self, arg):
        """Открывает карточку продукта и передает первым аргументом False (если функция вызвана по нажатию PushButton),
        либо список всех значений выбранной строки (если функция вызвана двойным нажатием по TableWidget)."""
        if arg is not False:
            arg = self.get_string_values(arg)
        product_card_window = AddProductDialog(arg, self.manager_id)
        product_card_window.exec_()
        self.insert_data_into_table()

    def count_quantity_products(self):
        """Формирует данные для заполнения нулевого склада"""
        resp = self.con.execute('''SELECT product_id, product_name, sku, unit_of_measurement,
                                      SUM (quantity) AS total_quantity, 
                                      MAX (expiration_date) AS exp_date
                                     FROM Products
                                     JOIN Stock
                                       ON Products.id = Stock.product_id
                                    GROUP BY product_id
                                    ORDER BY exp_date
                                    ''')
        data = resp.fetchall()
        columns = [el[0] for el in resp.description]
        return columns, data

    def open_transaction_card(self, arg):
        """Открывает карточку операции и передает первым аргументом False (если функция вызвана по нажатию PushButton),
        либо список всех значений выбранной строки (если функция вызвана двойным нажатием по TableWidget)."""
        self.data_from_outside = []
        self.dialog = DialogWindow()
        self.dialog.submitted.connect(self.update_data)
        self.dialog.show()

    @QtCore.pyqtSlot(list)
    def update_data(self, data):
        operation = DataProcessing(data)
        result = 'Error'
        if data[0] in ['sale', 'write-off']:
            result = operation.remove_from_warehouse()
        elif data[0] == 'acceptance':
            result = operation.add_to_warehouse()
        elif data[0] == 'movement':
            result = operation.movement_from_warehouse()
        self.announcement.setText(result)
        self.announcement.setStandardButtons(QMessageBox.Ok)
        self.announcement.exec_()
        self.insert_data_into_table()

    def add_word_doc(self):
        from docx import Document
        doc = Document()
        doc.add_heading("Вот доступные поля для создания шаблона", level=1)
        for k, v in self.pattern_doc.items():
            doc.add_paragraph('')
            doc.add_paragraph(f'{k}  -- {", ".join(v)}')
            doc.add_paragraph('')
        doc.save('My/shablon.docx')
        if platform.system() == 'Darwin':
            subprocess.call(('open', 'My/shablon.docx'))  # macOS
        elif platform.system() == 'Windows':
            os.startfile('My/shablon.docx')  # Windows
        else:
            subprocess.call(('xdg-open', 'My/shablon.docx'))  # linux variants

    def open_managers_list(self):
        """Открывает список менеджеров."""
        managers_window = Ui_listManager()
        managers_window.exec_()


if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)
    ui = MainDialog(file_db='../My/store_database.db')
    ui.show()
    sys.exit(app.exec_())
